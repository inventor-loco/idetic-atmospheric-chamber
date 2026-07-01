from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('Práctica de Laboratorio: IoT - MQTT con ESP32 y Sensor de CO2', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('Máster en Green Tech - IDeTIC', 1)
    
    doc.add_heading('1. Objetivos de la Práctica', level=2)
    doc.add_paragraph('• Configurar y programar un microcontrolador ESP32 utilizando el entorno de desarrollo Arduino IDE o PlatformIO.')
    doc.add_paragraph('• Leer datos analógicos/digitales de un sensor genérico de CO2.')
    doc.add_paragraph('• Establecer una conexión a una red Wi-Fi.')
    doc.add_paragraph('• Implementar el protocolo MQTT para publicar los datos del sensor en un broker.')
    doc.add_paragraph('• Fomentar el uso de tecnologías IoT aplicadas a la monitorización ambiental (Green Tech).')
    
    doc.add_heading('2. Materiales Requeridos', level=2)
    doc.add_paragraph('• 1x Placa de desarrollo ESP32 (e.g., ESP32-WROOM-32).')
    doc.add_paragraph('• 1x Sensor genérico de CO2 (e.g., MQ-135 o equivalente analógico/digital).')
    doc.add_paragraph('• Cables puente (jumper wires).')
    doc.add_paragraph('• Protoboard.')
    doc.add_paragraph('• Cable Micro-USB o USB-C para conexión al PC.')
    doc.add_paragraph('• PC con Arduino IDE o VS Code con PlatformIO instalado.')
    
    doc.add_heading('3. Marco Teórico Breve', level=2)
    p = doc.add_paragraph()
    p.add_run('MQTT (Message Queuing Telemetry Transport): ').bold = True
    p.add_run('Es un protocolo de mensajería ligero basado en el modelo publicación/suscripción. Es ideal para conexiones con dispositivos remotos donde el ancho de banda es premium y se requiere eficiencia (muy utilizado en IoT).\\n\\n')
    p.add_run('ESP32: ').bold = True
    p.add_run('Un microcontrolador de bajo costo y bajo consumo de energía con Wi-Fi y Bluetooth integrados, ideal para proyectos de IoT y Green Tech.')
    
    doc.add_heading('4. Esquema de Conexión', level=2)
    doc.add_paragraph('1. Conecta el pin VCC (o 5V/3.3V, según tu sensor) del sensor de CO2 al pin de 5V o 3.3V del ESP32.')
    doc.add_paragraph('2. Conecta el pin GND del sensor al pin GND del ESP32.')
    doc.add_paragraph('3. Conecta la salida analógica (A0) del sensor al pin ADC correspondiente en el ESP32 (por ejemplo, el Pin 34).')
    
    doc.add_heading('5. Desarrollo de la Práctica', level=2)
    
    doc.add_heading('Paso 5.1: Preparación del Entorno', level=3)
    doc.add_paragraph('Asegúrate de tener instaladas las bibliotecas "PubSubClient" (para MQTT) y "WiFi" en tu entorno de desarrollo.')
    
    doc.add_heading('Paso 5.2: Código Base', level=3)
    doc.add_paragraph('A continuación, se presenta un esquema de código básico que debes completar y cargar en tu ESP32:')
    
    code = """#include <WiFi.h>
#include <PubSubClient.h>

// Credenciales Wi-Fi
const char* ssid = "TU_SSID";
const char* password = "TU_PASSWORD";

// Configuración MQTT
const char* mqtt_server = "broker.hivemq.com"; // Broker público de prueba
const int mqtt_port = 1883;
const char* topic_co2 = "idetic/greentech/co2";

// Pin del Sensor de CO2
const int sensorPin = 34; // Pin ADC del ESP32

WiFiClient espClient;
PubSubClient client(espClient);

void setup_wifi() {
  delay(10);
  Serial.println();
  Serial.print("Conectando a ");
  Serial.println(ssid);
  WiFi.begin(ssid, password);
  while (WiFi.status() != WL_CONNECTED) {
    delay(500);
    Serial.print(".");
  }
  Serial.println("\\nWiFi conectado.");
}

void reconnect() {
  while (!client.connected()) {
    Serial.print("Intentando conexión MQTT...");
    if (client.connect("ESP32Client_GreenTech")) {
      Serial.println("conectado");
    } else {
      Serial.print("falló, rc=");
      Serial.print(client.state());
      Serial.println(" intentando de nuevo en 5 segundos");
      delay(5000);
    }
  }
}

void setup() {
  Serial.begin(115200);
  setup_wifi();
  client.setServer(mqtt_server, mqtt_port);
}

void loop() {
  if (!client.connected()) {
    reconnect();
  }
  client.loop();

  int sensorValue = analogRead(sensorPin);
  // Transformar valor si es necesario, según datasheet del sensor.
  
  String payload = String(sensorValue);
  Serial.print("Publicando nivel de CO2: ");
  Serial.println(payload);
  
  client.publish(topic_co2, payload.c_str());
  
  delay(5000); // Esperar 5 segundos antes de la siguiente lectura
}
"""
    p_code = doc.add_paragraph(code)
    p_code.style = 'No Spacing'
    for run in p_code.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 128)

    doc.add_heading('Paso 5.3: Visualización de Datos', level=3)
    doc.add_paragraph('1. Descarga e instala un cliente MQTT, como MQTT Explorer o usa una herramienta web.')
    doc.add_paragraph('2. Conéctate al broker "broker.hivemq.com" en el puerto 1883.')
    doc.add_paragraph('3. Suscríbete al topic "idetic/greentech/co2".')
    doc.add_paragraph('4. Observa cómo llegan los datos publicados por el ESP32 cada 5 segundos.')
    
    doc.add_heading('6. Tareas Propuestas (Para entregar)', level=2)
    doc.add_paragraph('• Calibración: Modifica el código para aplicar una ecuación de conversión (basada en el datasheet del sensor) que transforme el valor analógico leído (0-4095) en partes por millón (ppm).')
    doc.add_paragraph('• Formato JSON: Cambia el formato de publicación a JSON, incluyendo al menos "sensor_id", "tipo" y "valor" (ej. {"sensor_id": "esp32_01", "tipo": "CO2", "valor": 450}).')
    doc.add_paragraph('• Deep Sleep (Opcional - Green Tech extra): Modifica el código para que el ESP32 publique un dato y entre en modo "Deep Sleep" durante 1 minuto para ahorrar energía antes de la siguiente lectura.')
    
    doc.add_heading('7. Conclusiones y Evaluación', level=2)
    doc.add_paragraph('Redacta un breve informe (máx. 1 página) detallando los problemas encontrados durante la conexión Wi-Fi/MQTT, cómo resolviste la conversión a ppm y el impacto del uso del modo de bajo consumo (si lo implementaste) en el contexto de la sostenibilidad tecnológica.')

    doc.save('Practica_Laboratorio_IoT_MQTT.docx')

if __name__ == '__main__':
    main()
